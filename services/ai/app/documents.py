import base64
import binascii
import re
from dataclasses import dataclass
from io import BytesIO
from zipfile import BadZipFile, ZipFile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from lxml.etree import XMLSyntaxError  # type: ignore[import-untyped]
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from app.contracts.offering_document import OfferingDocumentFile
from app.core.errors import AiServiceError

MAX_FILE_BYTES = 8_000_000
MAX_DOCUMENT_CHARS = 50_000
MAX_TOTAL_CHARS = 80_000
MAX_PDF_PAGES = 50
MAX_DOCX_ENTRIES = 1_000
MAX_DOCX_UNCOMPRESSED_BYTES = 32_000_000


@dataclass(frozen=True)
class ExtractedDocument:
    filename: str
    text: str
    truncated: bool


def extract_documents(files: list[OfferingDocumentFile]) -> list[ExtractedDocument]:
    documents: list[ExtractedDocument] = []
    remaining = MAX_TOTAL_CHARS

    for uploaded in files:
        data = decode_file(uploaded)
        text = extract_text(uploaded, data)
        normalized = normalize_text(text)
        if not normalized:
            raise document_error(
                "AI_DOCUMENT_UNREADABLE",
                "The document did not contain readable text",
            )

        allowed = min(MAX_DOCUMENT_CHARS, remaining)
        if allowed <= 0:
            break
        truncated = len(normalized) > allowed
        documents.append(
            ExtractedDocument(
                filename=uploaded.filename,
                text=normalized[:allowed],
                truncated=truncated,
            )
        )
        remaining -= min(len(normalized), allowed)

    if not documents:
        raise document_error(
            "AI_DOCUMENT_UNREADABLE",
            "The documents did not contain readable text",
        )

    return documents


def decode_file(uploaded: OfferingDocumentFile) -> bytes:
    try:
        data = base64.b64decode(uploaded.base64_data, validate=True)
    except (binascii.Error, ValueError):
        raise document_error("AI_DOCUMENT_INVALID", "The document encoding is invalid") from None

    if not data or len(data) > MAX_FILE_BYTES:
        raise document_error("AI_DOCUMENT_INVALID", "The document size is invalid")
    return data


def extract_text(uploaded: OfferingDocumentFile, data: bytes) -> str:
    try:
        if uploaded.mime_type == "text/plain":
            if b"\x00" in data:
                raise ValueError("binary text")
            return data.decode("utf-8-sig")
        if uploaded.mime_type == "application/pdf":
            return extract_pdf(data)
        return extract_docx(data)
    except AiServiceError:
        raise
    except (
        BadZipFile,
        KeyError,
        OSError,
        PackageNotFoundError,
        PdfReadError,
        UnicodeDecodeError,
        ValueError,
        XMLSyntaxError,
    ):
        raise document_error("AI_DOCUMENT_UNREADABLE", "The document could not be read") from None


def extract_pdf(data: bytes) -> str:
    if not data.startswith(b"%PDF-"):
        raise ValueError("invalid PDF signature")

    reader = PdfReader(BytesIO(data), strict=True)
    if reader.is_encrypted or len(reader.pages) > MAX_PDF_PAGES:
        raise document_error(
            "AI_DOCUMENT_UNSUPPORTED",
            "Encrypted or unusually long PDF documents are not supported",
        )

    return "\n\n".join((page.extract_text() or "") for page in reader.pages)


def extract_docx(data: bytes) -> str:
    if not data.startswith(b"PK"):
        raise ValueError("invalid DOCX signature")

    with ZipFile(BytesIO(data)) as archive:
        entries = archive.infolist()
        if (
            len(entries) > MAX_DOCX_ENTRIES
            or sum(entry.file_size for entry in entries) > MAX_DOCX_UNCOMPRESSED_BYTES
            or any(entry.flag_bits & 0x1 for entry in entries)
            or "word/document.xml" not in archive.namelist()
        ):
            raise document_error(
                "AI_DOCUMENT_UNSUPPORTED",
                "This Word document cannot be processed safely",
            )

    document = Document(BytesIO(data))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def normalize_text(value: str) -> str:
    value = value.replace("\r\n", "\n").replace("\r", "\n")
    value = re.sub(r"[\t\f\v]+", " ", value)
    value = re.sub(r" +", " ", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def document_error(code: str, message: str) -> AiServiceError:
    return AiServiceError(code=code, message=message, status_code=422, retryable=False)
