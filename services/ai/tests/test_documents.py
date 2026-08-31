import base64
from io import BytesIO
from typing import Literal

import pytest
from docx import Document
from pypdf import PdfWriter

from app.contracts.offering_document import OfferingDocumentFile
from app.core.errors import AiServiceError
from app.documents import extract_documents


def uploaded_file(
    filename: str,
    mime_type: Literal[
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "text/plain",
    ],
    data: bytes,
) -> OfferingDocumentFile:
    return OfferingDocumentFile(
        filename=filename,
        mime_type=mime_type,
        base64_data=base64.b64encode(data).decode("ascii"),
    )


def test_extracts_utf8_text_document() -> None:
    documents = extract_documents(
        [uploaded_file("offers.txt", "text/plain", "Espresso: Rich house blend".encode())]
    )

    assert documents[0].filename == "offers.txt"
    assert documents[0].text == "Espresso: Rich house blend"
    assert documents[0].truncated is False


def test_extracts_docx_paragraphs_and_tables() -> None:
    document = Document()
    document.add_paragraph("Coffee subscriptions")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Office plan"
    table.cell(0, 1).text = "Weekly delivery"
    stream = BytesIO()
    document.save(stream)

    documents = extract_documents(
        [
            uploaded_file(
                "catalog.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                stream.getvalue(),
            )
        ]
    )

    assert "Coffee subscriptions" in documents[0].text
    assert "Office plan | Weekly delivery" in documents[0].text


def test_rejects_scanned_or_empty_pdf_without_pretending_to_run_ocr() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    stream = BytesIO()
    writer.write(stream)

    with pytest.raises(AiServiceError) as raised:
        extract_documents([uploaded_file("scan.pdf", "application/pdf", stream.getvalue())])

    assert raised.value.code == "AI_DOCUMENT_UNREADABLE"


def test_rejects_binary_content_claiming_to_be_text() -> None:
    with pytest.raises(AiServiceError) as raised:
        extract_documents([uploaded_file("offers.txt", "text/plain", b"offer\x00binary")])

    assert raised.value.code == "AI_DOCUMENT_UNREADABLE"
